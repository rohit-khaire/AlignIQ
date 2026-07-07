import os
import json
import csv
import io
from fastapi.responses import Response
from fpdf import FPDF, XPos, YPos
import logging

logger = logging.getLogger(__name__)

def generate_export(report_path: str, format_type: str) -> tuple[bytes, str, str]:
    """
    Reads the report from report_path and generates the export in the requested format.
    Returns a tuple of (content_bytes, media_type, filename).
    """
    if not os.path.exists(report_path):
        raise FileNotFoundError("Report file not found. Please run an analysis first.")
        
    with open(report_path, "r", encoding="utf-8") as f:
        report_data = json.load(f)
        
    if format_type == "json":
        content = json.dumps(report_data, indent=2).encode("utf-8")
        return content, "application/json", "compliance_report.json"
        
    elif format_type == "csv":
        output = io.StringIO()
        writer = csv.writer(output)
        
        # Write header
        writer.writerow(["Policy ID", "Policy Title", "Category", "Overall Status", "Matched Requirement", "Requirement Status", "Risk Score", "Reasoning", "Recommended Action"])
        
        results = report_data if isinstance(report_data, list) else report_data.get("report", [])
        for res in results:
            policy_id = res.get("policy_id", "")
            title = res.get("policy_title", "")
            overall_status = res.get("overall_status", "")
            category = res.get("category", "")
            
            reqs = res.get("requirements", [])
            if not reqs:
                writer.writerow([policy_id, title, category, overall_status, "", "", "", ""])
            else:
                for req in reqs:
                    req_text = req.get("requirement_text", "")
                    req_status = req.get("status", "")
                    risk_score = req.get("risk_score", "")
                    reasoning = req.get("reasoning", "")
                    action = req.get("recommended_action", "")
                    writer.writerow([policy_id, title, category, overall_status, req_text, req_status, risk_score, reasoning, action])
                    
        content = output.getvalue().encode("utf-8")
        return content, "text/csv", "compliance_report.csv"
        
    elif format_type == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        # --- HEADER ---
        pdf.set_fill_color(15, 23, 42) # Dark Slate
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", 'B', 18)
        pdf.multi_cell(0, 15, "AlignIQ Analysis Report", align='C', fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(10)
        
        # Calculate summary manually
        results = report_data if isinstance(report_data, list) else report_data.get("report", [])
        total_reqs = sum(r.get("summary", {}).get("total_requirements", 0) for r in results)
        satisfied = sum(r.get("summary", {}).get("satisfied", 0) for r in results)
        partial = sum(r.get("summary", {}).get("partially_satisfied", 0) for r in results)
        overall_score = int(((satisfied + (partial * 0.5)) / total_reqs) * 100) if total_reqs > 0 else 0
        
        # --- SUMMARY SECTION ---
        pdf.set_text_color(15, 23, 42)
        pdf.set_font("Helvetica", 'B', 14)
        pdf.multi_cell(0, 10, "Executive Summary", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        
        pdf.set_fill_color(248, 250, 252) # Very light grey
        pdf.set_font("Helvetica", '', 11)
        summary_text = (f"Total Master Policies Analyzed: {len(results)}\n"
                        f"Total Individual Requirements Checked: {total_reqs}\n"
                        f"Fully Satisfied: {satisfied}  |  Partially Satisfied: {partial}  |  Not Satisfied: {total_reqs - satisfied - partial}")
        pdf.multi_cell(0, 8, summary_text, fill=True, border=1, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
        pdf.set_font("Helvetica", 'B', 12)
        if overall_score >= 80:
            pdf.set_text_color(22, 163, 74) # Green
        elif overall_score >= 50:
            pdf.set_text_color(234, 88, 12) # Orange
        else:
            pdf.set_text_color(220, 38, 38) # Red
            
        pdf.multi_cell(0, 10, f"Global Compliance Score: {overall_score}%", align='R', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.set_text_color(0, 0, 0)
        pdf.ln(5)
        
        # --- DETAILED FINDINGS ---
        pdf.set_text_color(15, 23, 42)
        pdf.set_font("Helvetica", 'B', 14)
        pdf.multi_cell(0, 10, "Detailed Policy Findings", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.line(pdf.get_x(), pdf.get_y(), pdf.get_x() + 190, pdf.get_y())
        pdf.ln(5)
        
        for res in results:
            pdf.set_fill_color(241, 245, 249) # Slate 50
            pdf.set_text_color(15, 23, 42)    # Slate 900
            pdf.set_font("Helvetica", 'B', 12)
            
            clean_title = str(res.get('policy_title', '')).encode('latin-1', 'replace').decode('latin-1').replace('\n', ' ')
            pdf.multi_cell(0, 10, f"  {res.get('policy_id', '')} - {clean_title}", fill=True, border='L', new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            pdf.set_font("Helvetica", 'I', 10)
            pdf.set_text_color(100, 116, 139) # Slate 500
            
            category = res.get('category', 'Uncategorized')
            pdf.multi_cell(0, 6, f"  Category: {category}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.multi_cell(0, 6, f"  Overall Status: {res.get('overall_status', '')} (Confidence: {res.get('confidence', '')}%)", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            # Print master policy details if available
            mp_details = res.get('master_policy_details')
            if mp_details:
                pdf.set_font("Helvetica", '', 9)
                details_text = f"  Dept: {mp_details.get('department', 'N/A')} | Risk: {mp_details.get('risk_level', 'N/A')} | Owner: {mp_details.get('owner', 'N/A')}"
                pdf.multi_cell(0, 5, details_text, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                
            pdf.ln(3)
            
            reqs = res.get("requirements", [])
            if reqs:
                for req in reqs:
                    pdf.set_text_color(15, 23, 42)
                    pdf.set_font("Helvetica", 'B', 10)
                    clean_req = str(req.get('requirement_text', '')).encode('latin-1', 'replace').decode('latin-1').replace('\n', ' ')
                    pdf.multi_cell(0, 6, f"    - {clean_req}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    
                    status = req.get('status', '')
                    if status == "Satisfied":
                        pdf.set_text_color(22, 163, 74)
                    elif status == "Partially Satisfied":
                        pdf.set_text_color(234, 88, 12)
                    else:
                        pdf.set_text_color(220, 38, 38)
                        
                    pdf.set_font("Helvetica", 'B', 10)
                    pdf.multi_cell(0, 6, f"      Status: {status}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    
                    risk_score = req.get('risk_score')
                    if risk_score is not None:
                        if risk_score >= 12:
                            pdf.set_text_color(220, 38, 38)
                        elif risk_score >= 6:
                            pdf.set_text_color(234, 88, 12)
                        elif risk_score >= 3:
                            pdf.set_text_color(202, 138, 4)
                        else:
                            pdf.set_text_color(22, 163, 74)
                        pdf.multi_cell(0, 6, f"      Risk Score: {risk_score}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    
                    pdf.set_text_color(71, 85, 105)
                    pdf.set_font("Helvetica", '', 10)
                    clean_reasoning = str(req.get('reasoning', '')).encode('latin-1', 'replace').decode('latin-1').replace('\n', ' ')
                    pdf.multi_cell(0, 6, f"      Reasoning: {clean_reasoning}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.ln(3)
                    
            next_actions = res.get("next_actions", [])
            if next_actions:
                pdf.ln(2)
                pdf.set_text_color(2, 132, 199) # Sky 600
                pdf.set_font("Helvetica", 'B', 10)
                pdf.multi_cell(0, 6, "    Recommended Action Plan:", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                
                pdf.set_text_color(15, 23, 42)
                pdf.set_font("Helvetica", '', 10)
                for action in next_actions:
                    clean_action = str(action).encode('latin-1', 'replace').decode('latin-1').replace('\n', ' ')
                    pdf.multi_cell(0, 6, f"      \x95 {clean_action}", new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            pdf.ln(5)
            
        content = bytes(pdf.output())
        return content, "application/pdf", "compliance_report.pdf"
        
    else:
        raise ValueError(f"Unsupported format: {format_type}")

def generate_remediated_export(remediated_path: str, format_type: str, company_name: str) -> tuple[bytes, str, str]:
    """
    Reads the remediated policies report and generates the export in the requested format (docx or pdf).
    """
    if not os.path.exists(remediated_path):
        raise FileNotFoundError("Remediated policies not found.")
        
    with open(remediated_path, "r", encoding="utf-8") as f:
        data = json.load(f)
        
    safe_company_name = "".join([c if c.isalnum() else "_" for c in company_name])
    if not safe_company_name:
        safe_company_name = "Company"
        
    filename_base = f"{safe_company_name}_Remediated_Policies"
        
    if format_type == "docx":
        import docx
        doc = docx.Document()
        doc.add_heading(f"{company_name} - Remediated Policies", 0)
        
        for cat in data.get("remediated_categories", []):
            doc.add_heading(cat.get("category", "Policy"), level=1)
            text = cat.get("remediated_text", "")
            # Basic cleanup of markdown for docx if present, or just add as paragraph
            for line in text.split('\n'):
                if line.startswith('## '):
                    doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    doc.add_heading(line[4:], level=3)
                elif line.strip() == '---':
                    pass
                else:
                    if line.strip():
                        doc.add_paragraph(line)
            doc.add_page_break()
            
        output = io.BytesIO()
        doc.save(output)
        content = output.getvalue()
        return content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document", f"{filename_base}.docx"
        
    elif format_type == "pdf":
        pdf = FPDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        
        pdf.set_fill_color(15, 23, 42)
        pdf.set_text_color(255, 255, 255)
        pdf.set_font("Helvetica", 'B', 16)
        
        # Clean company name for FPDF latin-1
        clean_company = str(company_name).encode('latin-1', 'replace').decode('latin-1')
        pdf.multi_cell(0, 15, f"{clean_company} - Remediated Policies", align='C', fill=True, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
        pdf.ln(10)
        
        pdf.set_text_color(0, 0, 0)
        
        for cat in data.get("remediated_categories", []):
            pdf.set_font("Helvetica", 'B', 14)
            cat_name = str(cat.get('category', 'Policy')).encode('latin-1', 'replace').decode('latin-1')
            pdf.multi_cell(0, 10, cat_name, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            pdf.ln(2)
            
            pdf.set_font("Helvetica", '', 11)
            text = cat.get("remediated_text", "")
            for line in text.split('\n'):
                if line.strip() == '---':
                    pdf.ln(5)
                elif line.startswith('## '):
                    pdf.set_font("Helvetica", 'B', 12)
                    pdf.multi_cell(0, 8, line[3:].encode('latin-1', 'replace').decode('latin-1'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.set_font("Helvetica", '', 11)
                elif line.startswith('### '):
                    pdf.set_font("Helvetica", 'B', 11)
                    pdf.multi_cell(0, 7, line[4:].encode('latin-1', 'replace').decode('latin-1'), new_x=XPos.LMARGIN, new_y=YPos.NEXT)
                    pdf.set_font("Helvetica", '', 11)
                elif line.strip():
                    clean_line = str(line).encode('latin-1', 'replace').decode('latin-1')
                    pdf.multi_cell(0, 6, clean_line, new_x=XPos.LMARGIN, new_y=YPos.NEXT)
            
            pdf.add_page()
            
        content = bytes(pdf.output())
        return content, "application/pdf", f"{filename_base}.pdf"
        
    else:
        raise ValueError(f"Unsupported remediated format: {format_type}")
